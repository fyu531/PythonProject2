import os
import re
import logging
from flask import Flask, request, jsonify, send_from_directory
from flask_cors import CORS
from werkzeug.utils import secure_filename
from collections import defaultdict
import pdfplumber
from docx import Document
import spacy
import chardet

# 配置日志
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# 初始化Flask应用
app = Flask(__name__)
CORS(app)
UPLOAD_FOLDER = 'uploads'
TXT_FOLDER = 'extracted_txt'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(TXT_FOLDER, exist_ok=True)

# 加载中文NLP模型
try:
    nlp = spacy.load("zh_core_web_sm")
except OSError:
    logger.info("下载中文语言模型...")
    os.system("python -m spacy download zh_core_web_sm")
    nlp = spacy.load("zh_core_web_sm")

ALLOWED_EXTENSIONS = {'pdf', 'doc', 'docx'}


def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def save_text_to_file(text, original_filename):
    try:
        txt_filename = os.path.splitext(original_filename)[0] + '.txt'
        txt_path = os.path.join(TXT_FOLDER, secure_filename(txt_filename))
        with open(txt_path, 'w', encoding='utf-8') as f:
            f.write(text)
        logger.info(f"文本已保存至: {txt_path}")
        return txt_path
    except Exception as e:
        logger.error(f"保存TXT失败: {str(e)}")
        return None


def detect_encoding(file_path):
    with open(file_path, 'rb') as f:
        raw_data = f.read(10000)
    result = chardet.detect(raw_data)
    return result['encoding'] or 'utf-8'


def extract_text_from_file(filepath, original_filename):
    try:
        file_ext = filepath.rsplit('.', 1)[1].lower()
        text = ""
        logger.info(f"开始提取文件: {filepath}（类型: {file_ext}）")

        if file_ext == 'pdf':
            with pdfplumber.open(filepath) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text() or ""
                    text += page_text + "\n"
            if not text.strip():
                raise Exception(
                    "未提取到文本内容，可能是【扫描版PDF】（图片格式）。"
                    "请将其转换为可编辑文本后重试（可使用在线OCR工具）。"
                )
            logger.info(f"PDF提取成功，文本长度: {len(text)}字符")

        elif file_ext == 'docx':
            doc = Document(filepath)
            for para in doc.paragraphs:
                text += para.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    text += row_text + "\n"
            for rel in doc.part.rels.values():
                if "textbox" in rel.target_ref:
                    try:
                        textbox_part = rel.target_part
                        for para in textbox_part.element.xpath('.//w:p'):
                            textbox_text = "".join([t.text for t in para.xpath('.//w:t')])
                            text += textbox_text + "\n"
                    except Exception as e:
                        logger.warning(f"提取文本框时警告: {str(e)}（可忽略）")
            if not text.strip():
                raise Exception("DOCX文件无有效文本（可能为空文件或格式异常）")
            logger.info(f"DOCX提取成功，文本长度: {len(text)}字符")

        elif file_ext == 'doc':
            try:
                import textract
                raw_data = textract.process(filepath)
                encoding = detect_encoding(filepath)
                text = raw_data.decode(encoding, errors='replace')
                if not text.strip():
                    raise Exception("DOC文件未提取到文本内容（可能损坏或格式异常）")
                logger.info(f"DOC提取成功（编码: {encoding}），文本长度: {len(text)}字符")
            except ImportError:
                raise Exception("处理DOC文件需要安装textract: pip install textract")
            except Exception as e:
                raise Exception(
                    f"DOC文件提取失败: {str(e)}。"
                    "可能原因：缺少antiword/catdoc工具（Windows用户需手动安装）。"
                    "建议：将DOC另存为DOCX后重试，兼容性更好。"
                )

        else:
            raise Exception(f"不支持的文件类型: {file_ext}（仅支持{ALLOWED_EXTENSIONS}）")

        save_text_to_file(text, original_filename)
        return text

    except Exception as e:
        logger.error(f"提取失败: {str(e)}")
        raise


def detect_email(text):
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    matches = re.findall(email_pattern, text)
    return matches if matches else None


def detect_phone(text):
    phone_patterns = [
        r'1[3-9]\d{9}',
        r'0\d{2,3}-\d{7,8}',
        r'\(\d{3,4}\)\d{7,8}',
        r'\d{3,4}-\d{7,8}'
    ]
    for pattern in phone_patterns:
        matches = re.findall(pattern, text)
        if matches:
            return matches
    return None


def detect_major(text):
    major_patterns = [
        r'专业[:：]\s*([^\n]+)',
        r'所学专业[:：]\s*([^\n]+)',
        r'专业名称[:：]\s*([^\n]+)',
        r'毕业专业[:：]\s*([^\n]+)'
    ]
    for pattern in major_patterns:
        match = re.search(pattern, text)
        if match:
            return match.group(1).strip()
    education_patterns = [r'毕业于[:：]\s*([^\n]+)', r'学历[:：]\s*([^\n]+)']
    for pattern in education_patterns:
        match = re.search(pattern, text)
        if match:
            edu_info = match.group(1)
            major_match = re.search(r'[本科|硕士|博士][:：]?\s*([^\n]+)', edu_info)
            if major_match:
                return major_match.group(1).strip()
    return None


def detect_project_experience(text):
    project_keywords = ['项目经历', '项目经验', '参与项目', '负责项目', 'project experience', 'projects']
    for keyword in project_keywords:
        if re.search(rf'{re.escape(keyword)}', text, re.IGNORECASE):
            project_section = extract_section(text, keyword)
            return {'has_project': True, 'section_preview': project_section[:200] + '...' if project_section else None}
    return {'has_project': False}


def extract_section(text, keyword):
    match = re.search(rf'{re.escape(keyword)}', text, re.IGNORECASE)
    if not match:
        return None
    start_idx = match.end()
    end_pattern = re.compile(r'[\n\r]\s*[0-9一二三四五六七八九十]+[、.:：]')
    end_match = end_pattern.search(text, start_idx)
    if end_match:
        return text[start_idx:end_match.start()].strip()
    else:
        return text[start_idx:start_idx + 2000].strip()


SKILLS_DATABASE = {
    # 编程语言：根据傅泉智简历中"熟悉C/C++、Python"及编程场景调整
    '编程语言': {
        'Python': 8,  # 傅泉智用于LSTM模型开发，高频使用
        'C/C++': 8,  # 傅泉智主修课程及项目使用
        'Java': 5,
        'JavaScript': 5,
        'Go': 5,
        'Ruby': 4,
        'PHP': 4,
        'SQL': 7,  # 傅泉智涉及数据库编写
    },

    # 数据科学：结合傅泉智的机器学习项目和曾梓晟的数据分析工作
    '数据科学': {
        '机器学习': 10,  # 傅泉智LSTM模型核心技能
        '深度学习': 8,
        '数据分析': 9,  # 曾梓晟实习高频涉及，傅泉智项目需基础分析
        '金融数据分析': 10,  # 曾梓晟战略发展部实习核心技能
        '自然语言处理': 6,
        '计算机视觉': 9,  # 傅泉智涉及图像处理
        'TensorFlow': 6,
        'PyTorch': 6,
        'Scikit-learn': 7,
        'Pandas': 7,
        'NumPy': 7,
        '数据结构': 9,  # 傅泉智主修课程，项目核心基础
        '计算机组成': 8,
        '计算机网络': 7,
        '前端开发': 7  # 傅泉智掌握基础
    },

    # Web开发：基于傅泉智的前端基础调整
    'Web开发': {
        'React': 6,
        'Vue': 6,
        'Angular': 5,
        'Node.js': 5,
        'Express': 5,
        'Django': 6,
        'Flask': 6,
        'HTML': 5,
        'CSS': 5,
        '前端开发': 8,  # 傅泉智明确提及"基础的前端开发"
        '计算机组成': 7
    },

    # 工具和技术：补充两份简历中提及的专业工具
    '工具和技术': {
        'Git': 6,
        'Docker': 5,
        'Kubernetes': 5,
        'AWS': 4,
        'Azure': 4,
        'GCP': 4,
        'Linux': 6,
        'Jira': 5,
        'Jenkins': 5,
        'YOLO算法': 20,  # 傅泉智明确使用的核心算法
        'OpenCV': 15,  # 傅泉智图像处理项目核心工具
        'Arduino': 12,  # 傅泉智智能浇花系统核心工具
        'Office': 9,  # 曾梓晟"熟练使用"，高频办公场景
        'Wind金融终端': 12,  # 曾梓晟"精通"，战略分析核心工具
        'Choice金融终端': 12,  # 曾梓晟"精通"，金融数据核心工具
    },

    # 软技能：结合两份简历的实践经历调整
    '软技能': {
        '团队协作': 8,  # 曾梓晟社团/实习多次体现，傅泉智项目隐含
        '沟通能力': 8,  # 曾梓晟谈判合作、傅泉智项目协作均需
        '项目管理': 9,  # 曾梓晟主导多个项目，傅泉智为项目负责人
        '问题解决': 9,  # 傅泉智"善于解决问题"，曾梓晟优化流程体现
        '领导力': 8,  # 曾梓晟担任部长，傅泉智为项目负责人
        '创新能力': 7,
        '勤奋好学': 9,  # 傅泉智自我评价突出
        '踏实肯干': 10,  # 傅泉智自我评价核心优势
        '行业研究': 10,  # 曾梓晟深度解析多个行业，产出研究报告
        '营销策划': 8,  # 曾梓晟社团营销部工作核心技能
        '数据整合': 9  # 曾梓晟搭建数据框架、整合信息的核心能力
    }
}

EDUCATION_SCORES = {'博士': 10, '硕士': 8, '本科': 6, '大专': 4, '高中': 2}


def analyze_text(text):
    if not text or isinstance(text, str) and text.startswith("无法处理"):
        return {
            'is_valid_resume': False,
            'missing_info': ['无法提取简历内容'],
            'skills': {},
            'experience': 0,
            'education': '未找到教育信息',
            'score': 0,
            'comments': ['无法提取简历内容，请检查文件格式是否正确']
        }

    original_text = text
    text = text.lower()  # 转为小写统一匹配

    email = detect_email(original_text)
    phone = detect_phone(original_text)
    major = detect_major(original_text)
    project_exp = detect_project_experience(original_text)

    missing_info = []
    if not email:
        missing_info.append("邮箱")
    if not phone:
        missing_info.append("电话")
    if not major:
        missing_info.append("专业")
    if not project_exp['has_project']:
        missing_info.append("项目经历")

    is_valid_resume = len(missing_info) < 3

    # --------------------------
    # 核心优化：中文技能词匹配算法
    # --------------------------
    detected_skills = defaultdict(list)
    total_skill_score = 0
    max_possible_skill_score = sum(sum(category.values()) for category in SKILLS_DATABASE.values())

    # 技能词边界检查的例外字符列表（助词、标点等不影响技能词识别的字符）
    BOUNDARY_EXCEPTIONS = {'的', '了', '在', '和', '与', '或', '中', '，', '。', '、', '：', '(', ')', '（', '）'}

    # 将文本转为字符列表
    text_chars = list(text)
    text_length = len(text_chars)

    for category, skills in SKILLS_DATABASE.items():
        for skill, weight in skills.items():
            skill_lower = skill.lower()
            skill_chars = list(skill_lower)
            skill_length = len(skill_chars)

            if skill_length == 0:
                continue

            for i in range(text_length - skill_length + 1):
                # 逐个字符比对
                match = True
                for j in range(skill_length):
                    if text_chars[i + j] != skill_chars[j]:
                        match = False
                        break

                if match:
                    # 边界验证：优化中文语境下的边界判断
                    valid_front = True
                    valid_back = True

                    # 前边界验证
                    if i > 0:
                        prev_char = text_chars[i - 1]
                        # 如果前一个字符是中文字符且不在例外列表中，则无效
                        if '\u4e00' <= prev_char <= '\u9fa5' and prev_char not in BOUNDARY_EXCEPTIONS:
                            valid_front = False

                    # 后边界验证
                    if i + skill_length < text_length:
                        next_char = text_chars[i + skill_length]
                        # 如果后一个字符是中文字符且不在例外列表中，则无效
                        if '\u4e00' <= next_char <= '\u9fa5' and next_char not in BOUNDARY_EXCEPTIONS:
                            valid_back = False

                    if valid_front and valid_back:
                        detected_skills[category].append(skill)
                        total_skill_score += weight
                        logger.info(f"匹配到技能: {skill}（位置: {i}-{i + skill_length}）")
                        break  # 避免重复匹配
    # --------------------------
    # 后续逻辑保持不变
    # --------------------------
    skill_match_percentage = (total_skill_score / max_possible_skill_score) * 100 if max_possible_skill_score > 0 else 0

    education = '未找到教育信息'
    education_score = 0
    for edu_level, score in EDUCATION_SCORES.items():
        if edu_level in text:
            education = edu_level
            education_score = score
            break

    experience_years = 0
    experience_patterns = [
        r'(\d+)\s*[年載]工作经验', r'(\d+)\s*[年載]经验', r'工作\s*(\d+)\s*[年載]',
        r'有\s*(\d+)\s*[年載]', r'(\d+)\s*[年載]\s*工作'
    ]
    for pattern in experience_patterns:
        match = re.search(pattern, text)
        if match:
            try:
                experience_years = int(match.group(1))
                break
            except ValueError:
                continue

    total_score = (skill_match_percentage * 0.5) + (education_score * 2.5) + (min(experience_years, 10) * 2.5)

    comments = []
    if not is_valid_resume:
        comments.append(f"您的简历缺少关键信息：{', '.join(missing_info)}。完整的简历应包含这些信息以提高求职成功率。")
    else:
        if missing_info:
            comments.append(f"您的简历基本完整，但建议补充：{', '.join(missing_info)}，使简历更加完善。")
        else:
            comments.append("您的简历包含了所有关键信息，格式完整，给招聘方留下专业印象。")

    if not detected_skills:
        comments.append("简历中未明确提及相关技术技能，建议添加您掌握的专业技能。")
    else:
        top_skills = [skill for category in detected_skills.values() for skill in category][:3]
        comment = f"您的技能优势在于: {', '.join(top_skills)}。"
        if skill_match_percentage < 30:
            comment += " 建议进一步扩展您的技能栈，以增加竞争力。"
        elif skill_match_percentage < 60:
            comment += " 您具备一定的技能基础，但可以考虑学习更多前沿技术。"
        else:
            comment += " 您的技能组合非常出色，与行业需求高度匹配。"
        comments.append(comment)

    if education == '未找到教育信息':
        comments.append("简历中未明确提及教育背景信息，建议补充您的学历和专业。")
    else:
        if education_score >= 8:
            comments.append(f"您的{education}学历为您的职业发展提供了坚实的理论基础。")
        elif education_score >= 6:
            comments.append(f"本科学历是行业的普遍要求，您可以考虑进一步深造以提升竞争力。")
        else:
            comments.append(f"您的教育背景可能需要通过丰富的工作经验或专业技能来弥补。")

    if experience_years == 0:
        comments.append("简历中未明确提及工作经验，建议详细描述您的实习经历或项目经验。")
    elif experience_years < 3:
        comments.append(f"您拥有{experience_years}年工作经验，处于职业发展初期，建议专注于特定领域深入发展。")
    elif experience_years < 5:
        comments.append(f"您拥有{experience_years}年工作经验，已积累了一定的行业经验，可以考虑承担更多领导职责。")
    else:
        comments.append(f"您拥有{experience_years}年丰富的工作经验，在行业内已具备较强的竞争力。")

    if total_score < 40:
        comments.append("综合评分中等，建议在技能学习、学历提升或工作经验积累方面有所突破。")
    elif total_score < 70:
        comments.append("综合评分良好，您具备较强的竞争力，但仍有提升空间。")
    else:
        comments.append("综合评分优秀，您的简历在市场上具有很强的竞争力！")

    return {
        'is_valid_resume': is_valid_resume,
        'missing_info': missing_info,
        'contact_info': {'email': email, 'phone': phone},
        'major': major,
        'project_experience': project_exp,
        'skills': detected_skills,
        'experience': experience_years,
        'education': education,
        'score': round(total_score, 1),
        'comments': comments
    }


@app.route('/analyze', methods=['POST'])
def handle_analysis():
    if 'resume' not in request.files:
        return jsonify({'error': '未上传文件'}), 400

    file = request.files['resume']
    if file.filename == '':
        return jsonify({'error': '文件名无效'}), 400

    if not allowed_file(file.filename):
        return jsonify({'error': '文件类型不支持，仅支持: ' + ', '.join(ALLOWED_EXTENSIONS)}), 400

    filename = secure_filename(file.filename)
    filepath = os.path.join(UPLOAD_FOLDER, filename)
    file.save(filepath)

    try:
        text = extract_text_from_file(filepath, filename)
        txt_filename = os.path.splitext(filename)[0] + '.txt'
        txt_path = os.path.join(TXT_FOLDER, secure_filename(txt_filename))
        if os.path.exists(txt_path) and os.path.getsize(txt_path) > 0:
            with open(txt_path, 'r', encoding='utf-8', errors='replace') as f:
                text = f.read()
            logger.info(f"使用TXT文件进行分析: {txt_path}")

        result = analyze_text(text)
        return jsonify(result)
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        if os.path.exists(filepath):
            os.remove(filepath)
            logger.info(f"已删除临时文件: {filepath}")


@app.route('/')
def index():
    return send_from_directory('.', 'index.html')


# if __name__ == '__main__':
#     logger.info("""
#     首次运行请确保已安装依赖：
#     pip install flask flask-cors pdfplumber python-docx spacy textract chardet
#     python -m spacy download zh_core_web_sm
#     DOC文件处理需额外工具（如antiword），建议优先使用PDF或DOCX格式。
#     """)
#     # 关键修改：使用环境变量中的端口（Render随机分配）
#     port = int(os.environ.get('PORT', 5000))
#     app.run(host='0.0.0.0', port=port)